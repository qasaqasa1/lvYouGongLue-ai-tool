from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_COLOR_INDEX
from models import OutlineNode
import os
import re

def create_docx(location: str, title: str, content: str, filename: str = None) -> str:
    document = Document()
    
    # Title (H1 style)
    heading = document.add_heading(title, 0)
    heading.style.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A) # Dark Blue
    
    # Content (Enhanced HTML/Markdown to Docx converter)
    
    lines = content.split('\n')
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        # H2 / Title in Content
        if "<h2>" in line or line.startswith("# "):
            clean_text = re.sub(r'<[^>]+>', '', line).replace("# ", "")
            h = document.add_heading(clean_text, level=1)
            run = h.runs[0]
            run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A) # Dark Blue
            run.font.bold = True
            run.font.size = Pt(24)
            
        # List Item
        elif "<li>" in line or line.startswith("- "):
            clean_text = re.sub(r'<[^>]+>', '', line).replace("- ", "")
            p = document.add_paragraph(style='List Bullet')
            _add_rich_text(p, clean_text)
            p.paragraph_format.line_spacing = 1.5
            
        # Standard Paragraph
        else:
            clean_text = re.sub(r'<[^>]+>', '', line)
            if clean_text:
                p = document.add_paragraph()
                _add_rich_text(p, line) # Pass original line to parse inner tags
                p.paragraph_format.line_spacing = 1.5
            
    # Save
    output_dir = "outputs"
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)

    if not filename:
        # Sanitize title for filename
        safe_title = re.sub(r'[\\/*?:"<>|]', "", title)
        filename = f"{location}_{safe_title}.docx"
    
    # Ensure extension
    if not filename.endswith('.docx'):
        filename += '.docx'
        
    filepath = os.path.join(output_dir, filename)
    document.save(filepath)
    return filepath

def _add_rich_text(paragraph, text):
    """
    Parses simple HTML-like tags within a string and adds runs to paragraph.
    Supported: <strong>, <b>, <i>, <span style="background-color: ...">
    Also handles Emoji naturally.
    """
    # Very basic parser: split by tags
    # This is a regex that splits but keeps the delimiters
    tokens = re.split(r'(<[^>]+>)', text)
    
    current_bold = False
    current_italic = False
    current_highlight = False
    
    for token in tokens:
        if not token: 
            continue
            
        if token == "<strong>" or token == "<b>":
            current_bold = True
        elif token == "</strong>" or token == "</b>":
            current_bold = False
        elif token == "<i>":
            current_italic = True
        elif token == "</i>":
            current_italic = False
        elif "background-color" in token: # <span style="background-color: ...">
            current_highlight = True
        elif token == "</span>" and current_highlight:
            current_highlight = False
        elif token.startswith("<"): # Ignore other tags
            continue
        else:
            # Text content
            run = paragraph.add_run(token)
            if current_bold:
                run.bold = True
            if current_italic:
                run.italic = True
            if current_highlight:
                run.font.highlight_color = WD_COLOR_INDEX.YELLOW
